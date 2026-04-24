"""Module 4: Document Generation Engine.

LLM-templating for trade docs. Structured inputs -> editable Markdown output,
with optional DOCX export.
"""
from __future__ import annotations

import io
import json
from typing import Any

from docx import Document as DocxDocument

from app.ai import get_llm

_INPUT_GUIDE = (
    "The inputs JSON may contain: "
    "`sender` (the writer: full_name, company_name, title, email, phone), "
    "`supplier` (recipient: name, country, commodity, email, website, type), "
    "`deal` (title, commodity, volume_mt, buy_price, sell_price, freight_estimate, "
    "incoterms, currency, structure), "
    "an `opportunity_disclosure` block ({stage, commodity, volume_disclosure, "
    "geo_disclosure, single_deal, evaluating_origins, hold}) that gives you "
    "STAGE-GATED safe-to-disclose values — when this block is present you MUST "
    "use volume_disclosure verbatim (never the raw deal.volume_mt) and "
    "geo_disclosure verbatim (never the raw destination_port at stage 1-3); "
    "and a `negotiation` block ({stage, stage_label, side, reveal, ask, hold, tactics, "
    "known_intel, previously_disclosed, market_reference, supplier_quote, "
    "last_supplier_response}) that encodes the game-theory state of the "
    "relationship — see :mod:`app.ai.negotiation_strategy`. "
    "Use sender.full_name in the signature; if sender.company_name is empty, "
    "OMIT the company-name line entirely — never write a placeholder like "
    "'(your company — set it in Profile)', '[Your Company]', or 'Atlas'. "
    "Address the supplier by supplier.name. Reference the actual commodity from "
    "supplier.commodity or deal.commodity, never 'the specified commodity'. "
    "NEVER emit bracketed placeholders like '[Destination Country]', "
    "'[Specify time]', '[Your Name]' — if a value is missing, omit the "
    "line/clause entirely instead of leaving a placeholder."
)

_NEGOTIATION_RULES = (
    "NEGOTIATION STATE (hard constraints — non-negotiable): "
    "If inputs include a `negotiation` block, you MUST: "
    "(a) mention facts ONLY from negotiation.reveal (plus generic courtesy/signature); "
    "(b) NEVER mention any fact listed in negotiation.hold — these are leverage we are "
    "deliberately withholding; leaking one of them weakens the user's position; "
    "(c) include AT LEAST THREE explicit questions drawn from negotiation.ask, phrased "
    "as real questions (ending with '?') so the supplier must reply with data; "
    "(d) apply the tactics listed in negotiation.tactics — BATNA framing (hint at "
    "alternatives without bluffing), credible deadlines tied to an external event, "
    "anchoring discipline (never quote a price unless negotiation.stage >= 3), "
    "commitment escalation (ask for SCO / NCNDA / bank ref proportional to stage), "
    "and at least one non-price lever the counterparty can pull to defend margin; "
    "(e) if negotiation.previously_disclosed is non-empty, do not repeat those facts "
    "as if revealing them for the first time — reference them briefly or not at all; "
    "(f) if negotiation.known_intel contains their last quoted price/incoterms/MOQ, "
    "use them for targeting but DO NOT state them back as numbers we 'know' — this "
    "breaks information asymmetry."
)

DOC_SYSTEMS = {
    "outreach_email": (
        "You are a senior commodity trader writing a cold outreach email to a potential "
        "supplier. This is the FIRST message — the goal is to earn a reply, not to close "
        "a deal. "
        f"{_INPUT_GUIDE} "
        f"{_NEGOTIATION_RULES} "
        "STAGE-1 RULES (non-negotiable): "
        "(1) Do NOT quote any offer price, target price, bid price, or total deal value. "
        "Never invent a USD/MT or USD/lb number. Pricing is decided after the supplier "
        "responds — leading with a price weakens the buyer's position. "
        "(2) If inputs contain buy_price, sell_price or deal.volume_mt, treat them as "
        "INTERNAL ONLY and do not mention them. NEVER emit the exact target tonnage "
        "from deal.volume_mt or any specific MT figure — use opportunity_disclosure."
        "volume_disclosure verbatim instead (e.g. 'vessel-scale parcel "
        "(25,000-55,000 MT exploratory)'). The supplier MUST NOT be able to read the "
        "user's true target tonnage out of this email. "
        "(3) NEVER name the destination port. Use opportunity_disclosure.geo_disclosure "
        "verbatim — at stage 1 this is a region (e.g. 'West Africa'), never a country "
        "or port. "
        "(4) NEVER promise a first-shipment date or specific shipment window — say "
        "'subject to your origination cycle and our LC issuance' instead. "
        "(5) This is a SINGLE deal exploration, NOT a recurring contract. Do NOT use "
        "phrases like '12-month rolling contract', 'monthly programme', 'annual offtake'. "
        "Frame the volume as 'a single cargo' / 'one parcel' / 'this enquiry'. "
        "(6) Anchor on SPECS, not price: quality spec (e.g. ICUMSA 45 for sugar, grade "
        "#2 for corn), packaging (jumbo bags / 50kg / bulk), incoterms interest "
        "(FOB AND CFR — plural, to avoid anchoring), payment instrument family "
        "(DLC / SBLC) without tenor. "
        "(7) Ask the supplier for THEIR indicative pricing: BOTH FOB and CFR levels to "
        "the region, minimum order quantity, lead time from LC confirmation, payment "
        "terms accepted (SBLC, DLC at sight, etc). Phrase asks as REAL QUESTIONS "
        "ending with '?' — at least three '?' in the body. "
        "(8) Honest BATNA framing: when opportunity_disclosure.evaluating_origins is "
        "true, include ONE clause acknowledging multi-origin evaluation "
        "(e.g. 'we are comparing two origins this week' or 'evaluating multiple "
        "origins for this enquiry'). Do not bluff — keep it generic. "
        "(9) Signal seriousness without committing: mention readiness to execute NCNDA "
        "on interest, and that end-buyer LOI or proof-of-funds is available on request. "
        "(10) No filler, no hype, no emojis, no bracketed placeholders. Do not praise "
        "the supplier's reputation in more than one short clause. "
        "STRUCTURE (follow exactly): subject line that names commodity + "
        "opportunity_disclosure.volume_disclosure (the band, not a number) + "
        "opportunity_disclosure.geo_disclosure (region); one-line identification of "
        "the sender (use sender.full_name and, ONLY IF sender.company_name is non-empty, "
        "their company too — otherwise omit the company-name clause); one-line reason "
        "this supplier is relevant (reference their country / mill / product if known); "
        "a short line-separated spec block with the requirements from rule 6 (do NOT "
        "include a 'Destination port' line — geo_disclosure already covers it); the "
        "BATNA clause from rule 8; the explicit ask from rule 7 (≥3 '?' chars); the "
        "NCNDA readiness line from rule 9; soft close offering a 20-minute call this "
        "week. Under 200 words. Plain text (not Markdown). Sign off with the sender's "
        "full name, title (and company ONLY if non-empty), email, and phone on "
        "separate lines."
    ),
    "counter_offer_email": (
        "You are a senior commodity trader writing a counter-offer email in response to a "
        "supplier's quote. "
        f"{_INPUT_GUIDE} "
        "Additional context: inputs may include `market_reference` "
        "(exchange, ticker, price in USD/MT, source) and `supplier_quote` "
        "(supplier's offered price in USD/MT, incoterms, payment terms). "
        f"{_NEGOTIATION_RULES} "
        "STAGE-3 RULES: "
        "(1) This IS the email where a price appears — but it must be JUSTIFIED, not a "
        "random lowball. If market_reference is provided, explicitly anchor the counter "
        "to that futures price plus a transparent basis: 'based on ICE SB=F at "
        "$<price>/MT plus a $<basis>/MT freight-and-finance basis, our working level is "
        "$<counter>/MT CFR <port>'. "
        "(2) If market_reference is not provided, anchor to prevailing market context in "
        "general terms and request the supplier's basis rationale. "
        "(3) The counter should sit BELOW the midpoint between market reference and the "
        "supplier's quote — rule of thumb 3-5% below the supplier quote, adjusted by "
        "volume and payment terms. "
        "(4) Give the supplier two levers they can use to justify a smaller concession: "
        "willingness to pre-pay a deposit, or to accept DLC at sight on a shorter tenor. "
        "This protects the price without looking stubborn. "
        "(5) Never apologise for the counter. Never signal time pressure on the buyer "
        "side. Frame it as 'working level at which we can execute LC'. "
        "(6) Close with a concrete next step: either acceptance at the counter, or a "
        "counter-counter with specs the supplier can improve (tonnage upsize, tenor, "
        "ICUMSA grade). "
        "Under 200 words. Plain text. Sign off as in the outreach email."
    ),
    "follow_up_email": (
        "You are a senior commodity trader writing a STAGE-AWARE follow-up email to a "
        "supplier (or buyer) who has already responded to you. The negotiation block in "
        "the inputs tells you exactly what stage you are at (2-5) and what to do. "
        f"{_INPUT_GUIDE} "
        f"{_NEGOTIATION_RULES} "
        "Additional inputs for this doc type: `last_supplier_response` (the text of "
        "their most recent reply — read it carefully and address THEIR points, do not "
        "send a generic template). "
        "STAGE-SPECIFIC RULES: "
        "Stage 2 (first-response): acknowledge the SCO / indicative quote; ask for a "
        "FULL SCO on letterhead, payment-term RANGE, inspection-agency preference, and "
        "origination cycle; do NOT quote a price yourself; escalate to NCNDA signing. "
        "Stage 3 (counter-offer): this is the ONLY follow-up that states a price; anchor "
        "to market_reference + transparent basis (see counter_offer_email rules); give "
        "two non-price levers (deposit / shorter LC tenor) to defend their margin; close "
        "with concrete next-step (accept OR counter with specific terms). "
        "Stage 4 (terms-negotiation): price is agreed; negotiate packaging, freight "
        "split, inspection, LC format, tenor, loading window; now you may disclose "
        "destination port and deposit percentage; ask for bank reference and draft SPA. "
        "Stage 5 (close): all material terms settled; final disclosure is end-buyer "
        "identity under executed NCNDA; ask for draft SPA, LC pre-advice timing, and "
        "loading programme. "
        "STRUCTURE: Subject line references the commodity + stage context (e.g. 'RE: "
        "Brazil ICUMSA-45 — working level & next steps'); 1-2 sentence acknowledgement "
        "of what they sent; THE SPECIFIC asks for this stage; a credible deadline tied "
        "to an external event (shipment window, origination cycle, futures roll); a "
        "BATNA-aware sign-off that signals alternatives without bluffing. Under 220 "
        "words. Plain text. Sign off as in the outreach email."
    ),
    "spa_buyer": (
        "You are a commodity trade lawyer. Generate a complete, professional SALE AND "
        "PURCHASE AGREEMENT (SPA) from the buyer's perspective using the provided inputs. "
        "Output well-structured Markdown with numbered clauses. "
        f"{_INPUT_GUIDE} "
        "In the signature block, name the buyer as sender.company_name and the seller as "
        "supplier.name."
    ),
    "spa_supplier": (
        "You are a commodity trade lawyer. Generate a complete, professional SALE AND "
        "PURCHASE AGREEMENT (SPA) from the supplier/seller's perspective using the provided "
        "inputs. Output well-structured Markdown with numbered clauses. "
        f"{_INPUT_GUIDE} "
        "In the signature block, name the seller as sender.company_name and the buyer as "
        "supplier.name."
    ),
    "ncnda": (
        "You are a commodity trade lawyer. Draft a professional NON-CIRCUMVENTION, "
        "NON-DISCLOSURE AGREEMENT (NCNDA) using the inputs. Standard 3-year term unless "
        "overridden. Markdown output. "
        f"{_INPUT_GUIDE} "
        "Party A is sender.company_name, Party B is supplier.name. Include real "
        "signature blocks."
    ),
    "fpa": (
        "You are a commodity trade lawyer. Draft a FEE PROTECTION AGREEMENT (FPA) using "
        "the inputs. Markdown output. "
        f"{_INPUT_GUIDE}"
    ),
    "imfpa": (
        "You are a commodity trade lawyer. Draft an IRREVOCABLE MASTER FEE PROTECTION "
        "AGREEMENT (IMFPA) for intermediaries using the inputs. Markdown output. "
        f"{_INPUT_GUIDE}"
    ),
    "loi": (
        "You are a commodity trade lawyer. Draft a non-binding LETTER OF INTENT (LOI) "
        "from the buyer (sender.company_name) confirming readiness to purchase from "
        "supplier.name. Markdown output. "
        f"{_INPUT_GUIDE}"
    ),
}

DOC_TITLES = {
    "outreach_email": "Supplier Outreach Email",
    "counter_offer_email": "Counter-Offer Email",
    "follow_up_email": "Follow-Up Email",
    "spa_buyer": "Sale & Purchase Agreement (Buyer)",
    "spa_supplier": "Sale & Purchase Agreement (Supplier)",
    "ncnda": "Non-Circumvention, Non-Disclosure Agreement",
    "fpa": "Fee Protection Agreement",
    "imfpa": "Irrevocable Master Fee Protection Agreement",
    "loi": "Letter of Intent",
}


class DocumentGenerationService:
    def __init__(self) -> None:
        self.llm = get_llm()

    async def generate(self, doc_type: str, inputs: dict[str, Any]) -> tuple[str, str]:
        if doc_type not in DOC_SYSTEMS:
            raise ValueError(f"Unknown document type: {doc_type}")
        system = DOC_SYSTEMS[doc_type]
        user = (
            "Structured inputs (JSON):\n"
            f"{json.dumps(inputs, indent=2, default=str)}\n\n"
            f"Generate the {doc_type.replace('_', ' ')} document now."
        )
        content = await self.llm.complete(system, user, max_tokens=2000)
        title = DOC_TITLES[doc_type]
        return title, content.strip()

    def to_docx(self, title: str, markdown_content: str) -> bytes:
        """Basic Markdown -> DOCX conversion. Not feature-complete, but usable."""
        doc = DocxDocument()
        doc.add_heading(title, level=0)
        for line in markdown_content.splitlines():
            stripped = line.rstrip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
